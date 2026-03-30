import docx
import io

def extract_text_from_docx(file_input) -> str:
    """
    Extracts text directly from a DOCX file using python-docx.
    Accepts either a string file path or bytes.
    """
    text = []
    try:
        if isinstance(file_input, bytes):
            doc = docx.Document(io.BytesIO(file_input))
        else:
            doc = docx.Document(file_input)
            
        for para in doc.paragraphs:
            content = para.text.strip()
            if content:
                text.append(content)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        
    return "\n".join(text)
